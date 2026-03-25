import io
import base64
import logging
import hashlib
import json
import os
import sys
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from PIL import Image
import qrcode
from datetime import datetime
from typing import Optional

import database
from database import Report

logger = logging.getLogger("vbf")

# ─── Premium dokumentum stílusok ───
def _apply_premium_styles(doc, section, header_para):
    """Margók, alap betűtípus, címsorok és fejléc megjelenése."""
    # Margók: széles, professzionális
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.top_margin = Cm(1.5)
    section.bottom_margin = Cm(1.5)
    # Normal: Calibri 11pt, térköz után
    try:
        normal = doc.styles['Normal']
        normal.font.name = 'Calibri'
        normal.font.size = Pt(11)
        normal.paragraph_format.space_after = Pt(6)
        normal.paragraph_format.line_spacing = 1.15
    except Exception:
        pass
    # Heading 1: szakaszcímek
    try:
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Calibri'
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = RGBColor(31, 41, 55)
        h1.paragraph_format.space_before = Pt(16)
        h1.paragraph_format.space_after = Pt(8)
    except Exception:
        pass
    # Heading 2: alszakaszok
    try:
        h2 = doc.styles['Heading 2']
        h2.font.name = 'Calibri'
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.font.color.rgb = RGBColor(55, 65, 81)
        h2.paragraph_format.space_before = Pt(10)
        h2.paragraph_format.space_after = Pt(4)
    except Exception:
        pass
    # Fejléc: kisebb, szürke
    for run in header_para.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(107, 114, 128)
        run.font.name = 'Calibri'


def _style_table_header(table):
    """Tábla fejlécsora: szürke háttér, félkövér szöveg."""
    try:
        row0 = table.rows[0]
        for cell in row0.cells:
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), 'E2E8F0')
            shd.set(qn('w:val'), 'clear')
            tcPr.append(shd)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
    except Exception:
        pass


def generate_docx_stream(report: Report, db=None, share_url: Optional[str] = None, include_integrity: bool = False) -> io.BytesIO:
    doc = Document()
    
    settings = None
    if db and report.owner_id:
        owner = db.query(database.User).filter(database.User.id == report.owner_id).first()
        if owner and owner.company_id:
            settings = db.query(database.CompanySettings).filter(database.CompanySettings.company_id == owner.company_id).first()
        if not settings:
            settings = db.query(database.CompanySettings).filter(database.CompanySettings.owner_id == report.owner_id).first()
    
    # Header
    rep_type = report.report_type.upper() if report.report_type else "VBF"
    short_rep_type = "EPH" if rep_type == "EPH" else "VBF"
    
    if report.id and report.created_at:
        rep_id_str = f"{short_rep_type}-{report.created_at.year}-{report.id:03d}"
    else:
        rep_id_str = f"{short_rep_type}-TERVEZET"

    section = doc.sections[0]
    header = section.header
    if not header.paragraphs:
        header_para = header.add_paragraph()
    else:
        header_para = header.paragraphs[0]
        
    company_name = settings.company_name if settings and settings.company_name else "VBF Program"
    header_para.text = f"{company_name} | Azonosító: {rep_id_str}"
    header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    _apply_premium_styles(doc, section, header_para)

    # Insert Logo at the top if exists
    if settings and settings.logo_path and os.path.exists(settings.logo_path):
        try:
            # We add it to the first paragraph of the body
            logo_p = doc.add_paragraph()
            logo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            logo_img = Image.open(settings.logo_path)
            
            # Since WebP might have issues directly in Word via docx in some versions, convert temporarily to PNG in memory
            img_byte_arr = io.BytesIO()
            logo_img.save(img_byte_arr, format='PNG')
            img_byte_arr.seek(0)
            
            logo_p.add_run().add_picture(img_byte_arr, width=Cm(5))
        except Exception as e:
            logger.warning("Hiba a logó beillesztésekor: %s", e)
    
    # Title
    if rep_type == "VBF_IDOSZAKOS":
        title_text = "Időszakos Villamos Biztonsági Felülvizsgálati (VBF) Jegyzőkönyv"
    elif rep_type == "VBF_ELSO":
        title_text = "Első Villamos Biztonsági Felülvizsgálati (VBF) Jegyzőkönyv"
    elif rep_type == "VBF_BERBEADAS":
        title_text = "Bérbeadás Előtti Villamos Biztonsági Felülvizsgálati (VBF) Jegyzőkönyv"
    elif rep_type == "VBF_ELADAS":
        title_text = "Tulajdonosváltás / Eladás Előtti Villamos Biztonsági Felülvizsgálati (VBF) Jegyzőkönyv"
    elif rep_type == "EPH":
        title_text = "EPH Kialakítás és Bekötés Felülvizsgálati Jegyzőkönyv"
    else:
        title_text = f"Villamos Biztonsági Felülvizsgálati Jegyzőkönyv"

    # Címblokk – premium megjelenés
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title_text)
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = 'Calibri'
    title_run.font.color.rgb = RGBColor(15, 23, 42)
    title_para.paragraph_format.space_before = Pt(6)
    title_para.paragraph_format.space_after = Pt(4)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_para.add_run("Hivatalos Minősítő Irat")
    sub_run.font.size = Pt(11)
    sub_run.font.name = 'Calibri'
    sub_run.font.color.rgb = RGBColor(100, 116, 139)
    sub_run.italic = True
    subtitle_para.paragraph_format.space_after = Pt(12)

    # ── Összefoglaló kártya (1 oldal) ──
    c_data = report.client_data or {}
    d_data = report.defects_data or []
    critical_kw = ['életveszély', 'érintésvéd', 'pe vezető hiány', 'áramütés', 'tűzveszély', 'védővezető hiány', 'beégett', 'érinthető feszültség']
    serious_kw = ['szigetelés', 'rcd nem', 'ávk nem', 'fi-relé nem', 'hurokellenállás', 'túlterhelés', 'zárlat', 'hurokimpedancia', 'nem old ki']
    def _is_critical(desc):
        return any(kw in desc for kw in critical_kw)
    def _is_serious(desc):
        return any(kw in desc for kw in serious_kw)
    cnt_a = sum(1 for d in d_data if _is_critical((d.get('description') or '').lower()))
    cnt_b = sum(1 for d in d_data if _is_serious((d.get('description') or '').lower()) and not _is_critical((d.get('description') or '').lower()))
    r_val = c_data.get('reportResult', c_data.get('autoQualification', 'N/A'))
    
    card_p = doc.add_paragraph()
    card_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    card_run = card_p.add_run("ÖSSZEFOGLALÓ KÁRTYA")
    card_run.bold = True
    card_run.font.size = Pt(12)
    card_run.font.color.rgb = RGBColor(55, 65, 81)
    card_p.paragraph_format.space_after = Pt(8)
    
    card_table = doc.add_table(rows=6, cols=2)
    card_table.style = 'Table Grid'
    _style_table_header(card_table)
    rows = [
        ('Ügyfél / Megrendelő', c_data.get('customerName', 'N/A')),
        ('Vizsgálat helyszíne', c_data.get('siteAddress', 'N/A')),
        ('Vizsgálat dátuma', c_data.get('inspectionDate', report.created_at.strftime('%Y-%m-%d') if report.created_at else 'N/A')),
        ('Következő vizsgálat', c_data.get('nextInspectionDate', 'N/A')),
        ('Eredmény', r_val),
        ('Kritikus (A) / Súlyos (B) hibák', f'{cnt_a} / {cnt_b} db'),
    ]
    for i, (label, val) in enumerate(rows):
        card_table.rows[i].cells[0].text = label
        card_table.rows[i].cells[1].text = str(val)
    doc.add_paragraph()
    
    # ── QR-kód a borítólapra (megosztási / PDF link) ──
    if share_url:
        try:
            qr_cover = qrcode.QRCode(version=1, box_size=5, border=2, error_correction=qrcode.constants.ERROR_CORRECT_M)
            qr_cover.add_data(share_url)
            qr_cover.make(fit=True)
            qr_cover_img = qr_cover.make_image(fill_color="black", back_color="white")
            qr_cover_buf = io.BytesIO()
            qr_cover_img.save(qr_cover_buf, format='PNG')
            qr_cover_buf.seek(0)
            qr_cover_para = doc.add_paragraph()
            qr_cover_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qr_cover_para.add_run().add_picture(qr_cover_buf, width=Cm(3.5))
            qr_label_para = doc.add_paragraph("Megosztás / Letöltés")
            qr_label_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in qr_label_para.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(107, 114, 128)
        except Exception as e:
            doc.add_paragraph(f"[QR kód: {str(e)}]")
    
    doc.add_paragraph()
    subtitle_para.paragraph_format.space_after = Pt(6)
    
    # ── Tartalomjegyzék (TOC) ──
    doc.add_heading('Tartalomjegyzék', level=1)
    toc_para = doc.add_paragraph()
    toc_run = toc_para.add_run()
    
    # XML instruction for Word to insert TOC field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    toc_run._r.append(fldChar1)
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'TOC \\o "1-2" \\h \\z \\u'
    toc_run._r.append(instrText)
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    toc_run._r.append(fldChar2)
    
    # Placeholder text (Word updates this on first open if doc.settings allows)
    toc_run2 = toc_para.add_run('(A tartalomjegyzék frissítéséhez nyomjon Ctrl+A majd F9-et, vagy fogadja el a frissítést megnyitáskor)')
    toc_run2.font.color.rgb = RGBColor(128, 128, 128)
    toc_run2.font.italic = True
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    toc_run2._r.append(fldChar3)
    
    # Set Word to update fields on open
    doc.settings.element.append(OxmlElement('w:updateFields', {qn('w:val'): 'true'}))
    
    doc.add_page_break()
    
    # Preamble / Bevezető
    doc.add_heading('1. Cél és Vonatkozó Jogszabályok', level=1)
    if rep_type == "VBF_IDOSZAKOS":
        szabvany_ref = "• MSZ 10900 és MSZ HD 60364-6:2017 – Kisfeszültségű villamos berendezések időszakos ellenőrzése\n"
    elif rep_type == "VBF_ELSO":
        szabvany_ref = "• MSZ HD 60364-6:2017 – Kisfeszültségű villamos berendezések első ellenőrzése\n"
    else:
        szabvany_ref = "• MSZ HD 60364-6:2017 / MSZ 10900 – Érintésvédelmi felülvizsgálatok\n"
    
    # TvMI 7.7:2026.02.01 irányelv hivatkozás (hatályos 2026.02.01-től)
    tvmi_ref = "• TvMI 7.7:2026.02.01 – Villamos berendezések és villámvédelem tűzvédelmi felülvizsgálata (hatályos irányelv)\n"

    preamble_text = (
        "Jelen jegyzőkönyv a vizsgált villamos berendezés, villamos hálózat, illetve berendezések "
        "áramütés elleni védelmének, szabványos állapotának, valamint tűzvédelmi megfelelőségének "
        "minősítése céljából készült a Megrendelő megbízásából.\n\n"
        "A dokumentum az érvényben lévő nemzeti, illetve harmonizált európai szabványok szigorú betartása "
        "mellett lett összeállítva. A rögzített eredmények bírósági, hatósági, illetve biztosítói eljárások során "
        "hivatalos szakvéleményként / bizonyító erejű okiratként használhatók fel.\n\n"
        "1.1. A vizsgálat során alkalmazott főbb jogszabályok és rendeletek:\n"
        "• 40/2017. (XII. 4.) NGM rendelet – Az összekötő és felhasználói berendezésekről, valamint a Villamos Műszaki Biztonsági Szabályzatról (VMBSZ)\n"
        "• 54/2014. (XII. 5.) BM rendelet – Országos Tűzvédelmi Szabályzat (OTSZ)\n"
        f"{szabvany_ref}"
        f"{tvmi_ref}"
        "• MSZ EN 61140:2016 – Áramütés elleni védelem. A villamos berendezésekre és a villamos szerkezetekre vonatkozó közös szempontok\n"
        "• MSZ HD 60364-4-41:2018 – Biztonság. Áramütés elleni védelem\n\n"
        "1.2. A vizsgálat terjedelme, korlátai és a megrendelő felelőssége:\n"
        "A vizsgálat a helyszínen, a megrendelő által biztosított és bemutatott feltételek mellett történt. "
        "A vizsgálati eredmények és a kiadott minősítés kizárólag a vizsgált berendezésre, hozzáférhető hálózati pontokra "
        "és a vizsgálat időpontjában (a helyszíni eljárás napján) fennálló fizikai és méréstechnikai állapotra vonatkoznak. "
        "A felülvizsgáló nem vállal felelősséget a vizsgálat lezárását követően a hálózaton végrehajtott engedély nélküli "
        "beavatkozásokért, átalakításokért, bontásokért vagy a nem rendeltetésszerű használatból eredő meghibásodásokért, balesetekért."
    )
    p_intro = doc.add_paragraph(preamble_text)
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    # Client Data – Multi-vizsgálat: több helyszín egy jegyzőkönyvben
    doc.add_heading('2. Alapadatok és Helyszín', level=1)
    c_data = report.client_data or {}
    sites = c_data.get('sites') or []
    if isinstance(sites, list) and len(sites) > 1:
        for i, site in enumerate(sites, 1):
            if not isinstance(site, dict):
                continue
            doc.add_heading(f'2.{i} Helyszín: {site.get("name", site.get("siteName", f"Helyszín {i}"))}', level=2)
            p = doc.add_paragraph()
            p.add_run('Megrendelő / Üzemeltető: ').bold = True
            p.add_run(site.get('customerName', site.get('customer', c_data.get('customerName', 'N/A'))) + '\n')
            p.add_run('Vizsgálat helyszíne: ').bold = True
            p.add_run(site.get('siteAddress', site.get('address', 'N/A')) + '\n')
            p.add_run('Helyrajzi Szám (HRSZ): ').bold = True
            p.add_run(site.get('siteHrsz', site.get('hrsz', 'N/A')) + '\n')
            b_purpose = site.get('buildingPurpose', site.get('purpose', 'N/A'))
            b_otsz = site.get('buildingOtsz', site.get('otsz', ''))
            otsz_names = {'AK': 'Alacsony Kockázat', 'KK': 'Közepes Kockázat', 'MK': 'Magas Kockázat'}
            otsz_disp = f" [{otsz_names.get(b_otsz, b_otsz)}]" if b_otsz else ''
            p.add_run(f'Épület rendeltetése (OTSZ): {b_purpose}{otsz_disp}\n')
            if site.get('nextInspectionDate'):
                p.add_run('Következő vizsgálat: ').bold = True
                p.add_run(site.get('nextInspectionDate') + '\n')
    else:
        # Egyetlen helyszín (visszafelé kompatibilis)
        p = doc.add_paragraph()
        p.add_run('Megrendelő / Üzemeltető: ').bold = True
        p.add_run(c_data.get('customerName', 'N/A') + '\n')
        p.add_run('Vizsgálat helyszíne: ').bold = True
        p.add_run(c_data.get('siteAddress', 'N/A') + '\n')
        p.add_run('Helyrajzi Szám (HRSZ) / Azonosító: ').bold = True
        p.add_run(c_data.get('siteHrsz', 'N/A') + '\n')
        p.add_run('Épület rendeltetése (OTSZ): ').bold = True
        building_purpose = c_data.get('buildingPurpose', 'N/A')
        building_otsz = c_data.get('buildingOtsz', '')
        otsz_names = {'AK': 'Alacsony Kockázat', 'KK': 'Közepes Kockázat', 'MK': 'Magas Kockázat'}
        otsz_display = f" [{otsz_names.get(building_otsz, building_otsz)}]" if building_otsz else ''
        p.add_run(f"{building_purpose}{otsz_display}" + '\n')
        next_insp = c_data.get('nextInspectionDate', '')
        if next_insp:
            p.add_run('Következő kötelező felülvizsgálat (OTSZ): ').bold = True
            p.add_run(next_insp + '\n')
            
    # Mérés Körülményei & Alkalmazott Szabványok (Új)
    doc.add_heading('2.1 Mérés Körülményei és Jogszabályok', level=2)
    p_env = doc.add_paragraph()
    if c_data.get('envTemp') or c_data.get('envHumidity'):
        p_env.add_run('Hőmérséklet a vizsgálat idején: ').bold = True
        p_env.add_run(f"{c_data.get('envTemp', 'N/A')} °C, ")
        p_env.add_run('Páratartalom: ').bold = True
        p_env.add_run(f"{c_data.get('envHumidity', 'N/A')} %\n")
    
    p_env.add_run('Alkalmazott szabványok és jogszabályok: ').bold = True
    p_env.add_run(c_data.get('appliedStandards', 'MSZ HD 60364-6:2017') + '\n')
    
    if c_data.get('inspectionLimits'):
        p_env.add_run('Felülvizsgálat korlátai (mi nem lett vizsgálva): ').bold = True
        p_env.add_run(c_data.get('inspectionLimits') + '\n')
    
    # Inspector Data
    doc.add_heading('3. Felülvizsgáló és Műszerek', level=1)
    p2 = doc.add_paragraph("A vizsgálatot a vonatkozó jogszabályokban előírt szakmai végzettséggel és érvényes vizsgabizonyítvánnyal rendelkező személy végezte.\n")
    p2.add_run('Felülvizsgáló neve/cége: ').bold = True
    p2.add_run(c_data.get('inspectorName', 'N/A') + '\n')
    
    p2.add_run('Vizsgabizonyítvány száma: ').bold = True
    p2.add_run(c_data.get('inspectorLicense', 'N/A') + '\n')
    
    p2.add_run('Alkalmazott Mérőműszer: ').bold = True
    p2.add_run(c_data.get('instrumentType', 'N/A') + '\n')
    
    p2.add_run('Kalibrálás érvényessége: ').bold = True
    p2.add_run(c_data.get('instrumentCal', 'N/A') + '\n')
    
    if c_data.get('instrumentError'):
        p2.add_run('Műszer Mérési Bizonytalansága: ').bold = True
        p2.add_run(c_data.get('instrumentError') + '\n')

    # Visual Checklist
    section_num = 4
    visual = c_data.get('visualChecks', {})
    if isinstance(visual, dict) and visual:
        doc.add_heading(f'{section_num}. Szemrevételezéses Ellenőrzések (6.4.2 MSZ HD 60364-6)', level=1)
        v_table = doc.add_table(rows=1, cols=2)
        v_table.style = 'Table Grid'
        _style_table_header(v_table)
        v_hdr = v_table.rows[0].cells
        v_hdr[0].text = 'Ellenőrzött pont'
        v_hdr[1].text = 'Minősítés'
        
        checks = [
            ('Azonosító jelek, feliratok megléte', visual.get('id_marks', False)),
            ('Áramütés elleni védelem kialakítása', visual.get('protection', False)),
            ('Tűzvédelmi óvintézkedések', visual.get('fire', False)),
            ('Vezetők kiválasztása, terhelhetőség', visual.get('conduction', False)),
            ('Csatlakozások, kötések megfelelősége', visual.get('connection', False)),
            ('Karbantarthatóság, hozzáférhetőség', visual.get('access', False))
        ]
        for label, val in checks:
            v_row = v_table.add_row().cells
            v_row[0].text = label
            v_row[1].text = "Megfelelő" if val else "Nem felel meg / Nem vizsgált"
        doc.add_paragraph()
        section_num += 1

    # Single-line Diagram (Egyvonalas rajz)
    diagram_b64 = getattr(report, 'diagram_image', None)
    if diagram_b64 and diagram_b64.startswith('data:image'):
        doc.add_page_break()
        doc.add_heading(f'{section_num}. Egyvonalas Rajz (Áramkör Áttekintés)', level=1)
        doc.add_paragraph(
            'Az alábbi egyvonalas rajz a vizsgált villamos hálózat '
            'áramköri felépítését és az elosztók hierarchiáját mutatja be.'
        )
        try:
            # Handle data URL with or without metadata prefix
            if ',' in diagram_b64:
                b64_str = diagram_b64.split(',')[1]
            else:
                b64_str = diagram_b64
                
            img_bytes = io.BytesIO(base64.b64decode(b64_str))
            
            # Add image with center alignment
            diag_p = doc.add_paragraph()
            diag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            diag_run = diag_p.add_run()
            # Scaling to fit page width (16cm) while maintaining aspect ratio
            diag_run.add_picture(img_bytes, width=Cm(16))
            
            doc.add_paragraph() # Spacer
            section_num += 1
        except Exception as e:
            logger.warning("Hiba a diagram beillesztésekor: %s", e)
            doc.add_paragraph(f"[Hiba a rajz beillesztésekor: {str(e)}]")

    # EPH Specific Data — jelenjen meg VBF jegyzőkönyvekben is, ha ki van töltve
    has_eph_fields = any(
        c_data.get(key)
        for key in ["ephGasRequired", "ephGasMeter", "ephPenSep", "ephEarthMethod", "ephRaValue", "ephConductor"]
    )
    if rep_type == "EPH" or has_eph_fields:
        doc.add_heading(f'{section_num}. EPH és Földelés Specifikus Adatok', level=1)
        p3 = doc.add_paragraph()
        p3.add_run('Gázszolgáltató (Gázmérő) bevonása szükséges: ').bold = True
        p3.add_run(c_data.get('ephGasRequired', 'Nem') + '\n')
        
        p3.add_run('Gázmérő gyári száma: ').bold = True
        p3.add_run(c_data.get('ephGasMeter', 'N/A') + '\n')
        
        p3.add_run('PE-N Szétválasztás Helye: ').bold = True
        p3.add_run(c_data.get('ephPenSep', 'N/A') + '\n')
        
        p3.add_run('Földelési Ellenállás Mérési Módszer: ').bold = True
        p3.add_run(c_data.get('ephEarthMethod', 'N/A') + '\n')
        
        p3.add_run('Mért Földelési Ellenállás (Ra): ').bold = True
        if c_data.get('ephEarthNotMeasurable'):
            p3.add_run('A mérés a helyszín adottságai (pl. társasház, zárt mérési pont) miatt nem volt kivitelezhető (Rpe folytonosság mérve).\n')
        else:
            p3.add_run(str(c_data.get('ephRaValue', 'N/A')) + ' Ω\n')
        
        p3.add_run('EPH Fővezeték Keresztmetszete: ').bold = True
        p3.add_run(c_data.get('ephConductor', 'N/A') + ' mm²\n')
        
        if c_data.get('ephDeclaration'):
            doc.add_paragraph().add_run('EPH Nyilatkozat: Alulírott nyilatkozom, hogy az egyidejűleg érinthető idegen fémszerkezeteket megfelelően bekötötték a hálózatba.').italic = True

        section_num += 1
    
    # Measurements – measurements_data a DB-ben dict, nem lista
    meas_data = report.measurements_data if isinstance(report.measurements_data, dict) else {}
    if meas_data:
        doc.add_heading(f'{section_num}. Mérési Eredmények', level=1)
        sub_num = 1
        
        # VBF Measurements
        if rep_type.startswith("VBF") or rep_type == "VVF":
            # RPE
            rpe_list = meas_data.get('rpe', [])
            if rpe_list:
                doc.add_heading(f'{section_num}.{sub_num} Védővezető folytonosság (Rpe)', level=2)
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                _style_table_header(table)
                hdr = table.rows[0].cells
                hdr[0].text = 'Pont'
                hdr[1].text = 'Mérés Helye'
                hdr[2].text = 'Rpe [Ω]'
                hdr[3].text = 'Megfelel'
                for r in rpe_list:
                    row = table.add_row().cells
                    row[0].text = r.get('point', '')
                    row[1].text = r.get('loc', '')
                    row[2].text = r.get('val', '')
                    row[3].text = r.get('pass', '')
                sub_num += 1
                doc.add_paragraph()
            
            # Insulation
            ins_list = meas_data.get('insulation', [])
            if ins_list:
                doc.add_heading(f'{section_num}.{sub_num} Szigetelési ellenállás (500V DC)', level=2)
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                _style_table_header(table)
                hdr = table.rows[0].cells
                hdr[0].text = 'Áramkör'
                hdr[1].text = 'Riso L-N [MΩ]'
                hdr[2].text = 'Riso L-PE [MΩ]'
                hdr[3].text = 'Riso N-PE [MΩ]'
                hdr[4].text = 'Megfelel'
                for r in ins_list:
                    row = table.add_row().cells
                    row[0].text = r.get('circuit', '')
                    row[1].text = r.get('ln', '')
                    row[2].text = r.get('lpe', '')
                    row[3].text = r.get('npe', '')
                    row[4].text = r.get('pass', '')
                sub_num += 1
                doc.add_paragraph()
            
            # Loop
            loop_list = meas_data.get('loop', [])
            if loop_list:
                doc.add_heading(f'{section_num}.{sub_num} Hurokellenállás (Zs) és Hibavédelem', level=2)
                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                _style_table_header(table)
                hdr = table.rows[0].cells
                hdr[0].text = 'Áramkör / Pont'
                hdr[1].text = 'Kikapcsoló szerv'
                hdr[2].text = 'Hely (X/n)'
                hdr[3].text = 'Zs [Ω]'
                hdr[4].text = 'Megfelel'
                for r in loop_list:
                    row = table.add_row().cells
                    row[0].text = r.get('circuit', '')
                    row[1].text = r.get('device', '')
                    row[2].text = r.get('loc', '')
                    row[3].text = r.get('zs', '')
                    row[4].text = r.get('pass', '')
                sub_num += 1
                doc.add_paragraph()
            
            # RCD
            rcd_list = meas_data.get('rcd', [])
            if rcd_list:
                doc.add_heading(f'{section_num}.{sub_num} FI-relé (ÁVK) részletes vizsgálat', level=2)
                table = doc.add_table(rows=1, cols=9)
                table.style = 'Table Grid'
                _style_table_header(table)
                hdr = table.rows[0].cells
                hdr[0].text = 'Áramkör'
                hdr[1].text = 'Típus'
                hdr[2].text = 'IΔn [mA]'
                hdr[3].text = '0.5xIΔn'
                hdr[4].text = '1xIΔn [ms]'
                hdr[5].text = '5xIΔn [ms]'
                hdr[6].text = 'IΔ [mA]'
                hdr[7].text = 'Uc [V]'
                hdr[8].text = 'Megfelel'
                for r in rcd_list:
                    row = table.add_row().cells
                    row[0].text = str(r.get('circ', ''))
                    row[1].text = str(r.get('type', ''))
                    row[2].text = str(r.get('idn', ''))
                    row[3].text = str(r.get('test05', ''))
                    row[4].text = str(r.get('t1', ''))
                    row[5].text = str(r.get('t5', ''))
                    row[6].text = str(r.get('ramp', ''))
                    row[7].text = str(r.get('uc', ''))
                    row[8].text = str(r.get('pass', ''))
                sub_num += 1
                doc.add_paragraph()
                
            # Handheld Tools (M3)
            tools_list = meas_data.get('tools', [])
            if tools_list:
                doc.add_heading(f'{section_num}.{sub_num} Kéziszerszámok vizsgálata', level=2)
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                _style_table_header(table)
                hdr = table.rows[0].cells
                hdr[0].text = 'Megnevezés'
                hdr[1].text = 'Azonosító'
                hdr[2].text = 'Szig. ell. [MΩ]'
                hdr[3].text = 'Megfelel'
                for r in tools_list:
                    row = table.add_row().cells
                    row[0].text = r.get('name', '')
                    row[1].text = r.get('id', '')
                    row[2].text = r.get('val', '')
                    row[3].text = r.get('pass', '')
                sub_num += 1
                doc.add_paragraph()

            # SELV/PELV (M2)
            selv_list = meas_data.get('selv', [])
            if selv_list:
                doc.add_heading(f'{section_num}.{sub_num} SELV / PELV / Elválasztás', level=2)
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                _style_table_header(table)
                hdr = table.rows[0].cells
                hdr[0].text = 'Hely / Áttétel'
                hdr[1].text = 'V [V]'
                hdr[2].text = 'P-S [MΩ]'
                hdr[3].text = 'P-T [MΩ]'
                hdr[4].text = 'S-T [MΩ]'
                hdr[5].text = 'Megfelel'
                for r in selv_list:
                    row = table.add_row().cells
                    row[0].text = r.get('loc', '')
                    row[1].text = r.get('v', '')
                    row[2].text = r.get('ps', '')
                    row[3].text = r.get('pt', '')
                    row[4].text = r.get('st', '')
                    row[5].text = r.get('pass', '')
                sub_num += 1
                doc.add_paragraph()
                
        elif rep_type == "EPH":
            eph_list = meas_data.get('eph_cont', [])
            if eph_list:
                doc.add_heading(f'{section_num}.{sub_num} EPH Bekötések Folytonossága', level=2)
                table = doc.add_table(rows=1, cols=7)
                table.style = 'Table Grid'
                _style_table_header(table)
                hdr = table.rows[0].cells
                hdr[0].text = 'Sorszám'
                hdr[1].text = 'Bekötött Elem'
                hdr[2].text = 'Hely'
                hdr[3].text = 'Vezető (Keresztm.)'
                hdr[4].text = 'Kötés módja'
                hdr[5].text = 'Folytonosság [Ω]'
                hdr[6].text = 'Megfelel'
                for r in eph_list:
                    row = table.add_row().cells
                    row[0].text = str(r.get('idx', ''))
                    row[1].text = str(r.get('elem', ''))
                    row[2].text = str(r.get('loc', ''))
                    row[3].text = str(r.get('mat', ''))
                    row[4].text = str(r.get('conn', ''))
                    row[5].text = str(r.get('val', ''))
                    row[6].text = str(r.get('pass', ''))
                sub_num += 1
                doc.add_paragraph()

        # Print photos attached to measurement circuits
        meas_photos = []
        for m_key in ['rpe', 'insulation', 'loop', 'rcd', 'tools', 'selv', 'eph_cont']:
            items = meas_data.get(m_key, [])
            if not isinstance(items, list): continue
            for m_row in items:
                if not isinstance(m_row, dict): continue
                p_data = m_row.get('photo')
                if p_data and isinstance(p_data, str) and p_data.startswith('data:image'):
                    desc = "Mérés Kép"
                    if m_key == 'rpe': desc = f"Védővezető Rpe - {m_row.get('loc', '')}"
                    elif m_key == 'insulation': desc = f"Szigetelés - {m_row.get('circuit', '')}"
                    elif m_key == 'loop': desc = f"Hurokellenállás - {m_row.get('circuit', '')} ({m_row.get('loc', '')})"
                    elif m_key == 'rcd': desc = f"Fi-Relé - {m_row.get('circ', '')}"
                    elif m_key == 'tools': desc = f"Kéziszerszám - {m_row.get('name', '')}"
                    elif m_key == 'selv': desc = f"SELV - {m_row.get('loc', '')}"
                    elif m_key == 'eph_cont': desc = f"EPH - {m_row.get('elem', '')} ({m_row.get('loc', '')})"
                    meas_photos.append((desc, p_data))
                    
        if meas_photos:
            doc.add_heading(f'{section_num}. Mérési Áramkörökhöz / Sorokhoz csatolt fényképek', level=1)
            for desc, photo_data in meas_photos:
                try:
                    b64_str = photo_data.split(',')[1]
                    img_bytes = io.BytesIO(base64.b64decode(b64_str))
                    doc.add_paragraph(f"{desc}:").bold = True
                    doc.add_paragraph().add_run().add_picture(img_bytes, width=Cm(12))
                except Exception as e:
                    doc.add_paragraph(f"[Hiba a mérés képének beillesztésekor: {str(e)}]")
        section_num += 1
        doc.add_paragraph()

    # Defects - Szabvány szerinti kategorizálás
    doc.add_heading(f'{section_num}. Feltárt Hibák és Hiányosságok', level=1)
    
    # Severity legend
    sev_legend = doc.add_paragraph()
    sev_legend.add_run("Hibakategóriák jelentősége:\n").bold = True
    sev_legend.add_run("(A) Közvetlen élet- és tűzveszély  ")
    sev_legend.add_run("(B) Súlyos hiba (soron kívül javítandó)  ")
    sev_legend.add_run("(C) Karbantartási hiba  ")
    sev_legend.add_run("(D) Felújításkor javítandó\n")
    
    d_data = report.defects_data or []
    if not d_data:
        doc.add_paragraph("A vizsgálat során nem tártunk fel hibát vagy hiányosságot.")
    else:
        # Severity keyword detection (mirror of frontend logic)
        critical_kw = ['életveszély', 'érintésvéd', 'pe vezető hiány', 'áramütés', 
                       'tűzveszély', 'védővezető hiány', 'beégett', 'érinthető feszültség']
        serious_kw = ['szigetelés', 'rcd nem', 'ávk nem', 'fi-relé nem', 'hurokellenállás',
                      'túlterhelés', 'zárlat', 'hurokimpedancia', 'nem old ki']
        maintenance_kw = ['dobozfedél', 'fedél hiány', 'csatlakozó laza', 'felirat hiány',
                          'jelölés hiány', 'kötés laza', 'sorkapocs', 'burkolat sérült']
        renovation_kw = ['vezetékszínezés', 'régi szabvány', 'elavult', 'korszerűtlen',
                         'alumínium vezető', 'téves színezés', 'nullázás']
        
        for idx, defect in enumerate(d_data, 1):
            desc_lower = defect.get('description', '').lower()
            
            # Determine severity
            if any(kw in desc_lower for kw in critical_kw):
                severity = '(A)'
                sev_name = 'Közvetlen élet- és tűzveszély'
                sev_color = RGBColor(255, 0, 0)
            elif any(kw in desc_lower for kw in serious_kw):
                severity = '(B)'
                sev_name = 'Súlyos hiba (soron kívül javítandó)'
                sev_color = RGBColor(255, 80, 0)
            elif any(kw in desc_lower for kw in maintenance_kw):
                severity = '(C)'
                sev_name = 'Karbantartási hiba'
                sev_color = RGBColor(255, 140, 0)
            elif any(kw in desc_lower for kw in renovation_kw):
                severity = '(D)'
                sev_name = 'Felújításkor javítandó'
                sev_color = RGBColor(128, 128, 0)
            else:
                severity = '(C)'
                sev_name = 'Karbantartási hiba'
                sev_color = RGBColor(255, 140, 0)
            
            doc.add_heading(f"Hiba #{idx} — {severity} {sev_name}", level=2)
            dp = doc.add_paragraph()
            
            # Severity badge
            sev_run = dp.add_run(f"[{severity}] ")
            sev_run.bold = True
            sev_run.font.color.rgb = sev_color
            
            dp.add_run("Leírás és Javaslat:\n").bold = True
            dp.add_run(defect.get('description', 'N/A') + "\n\n")

            # Ajánlott javítási szöveg (MSZ alapján) – automatikus
            repair_text = defect.get('repairSuggestion', '')
            if not repair_text or not str(repair_text).strip():
                if any(kw in desc_lower for kw in ['védővezető', 'rpe', 'folytonosság', 'pe hiány']):
                    repair_text = 'A védővezető folytonosságának helyreállítása szükséges. Ellenőrizze a csatlakozások megfelelőségét, a PE vezeték épségét. MSZ HD 60364-6:2017 §61.3.2; VMBSZ 40/2017. (XII.4.) NGM rendelet.'
                elif any(kw in desc_lower for kw in ['szigetelés', 'riso', 'insulation']):
                    repair_text = 'Szigetelési ellenállás mérés és vizuális ellenőrzés. Serült vezeték cseréje, nedves/károsodott részek javítása. MSZ HD 60364-6:2017 §61.3.3; TvMI 7.7:2026.02.01.'
                elif any(kw in desc_lower for kw in ['hurok', 'hurokellenállás', 'hurokimpedancia', 'zs']):
                    repair_text = 'A hurokimpedancia csökkentése: hibavédelmi berendezés csere, vezetékkeresztmetszet növelés, csatlakozások ellenőrzése. MSZ HD 60364-4-41 §411.4.'
                elif any(kw in desc_lower for kw in ['rcd', 'ávk', 'fi-relé', 'kioldás', 'nem old']):
                    repair_text = 'FI-relé (ÁVK) cseréje vagy javítása. Ellenőrizze a bekötést és a kioldási időket. MSZ HD 60364-6:2017 §61.3.7; MSZ EN 61008-1.'
                elif any(kw in desc_lower for kw in ['tűz', 'égett', 'ív', 'beégett']):
                    repair_text = 'A károsodott berendezés azonnali cseréje. Tűzvédelmi szabályok betartása (OTSZ 54/2014. BM rendelet).'
                elif any(kw in desc_lower for kw in ['dobozfedél', 'fedél hiány', 'burkolat']):
                    repair_text = 'Hiányzó fedél vagy burkolat pótlása. MSZ HD 60364-6 szerinti szemrevételezés.'
                elif any(kw in desc_lower for kw in ['csatlakozó laza', 'kötés laza', 'sorkapocs']):
                    repair_text = 'Csatlakozások meghúzása, megfelelő szorítónyomat. Ellenőrizze a hőnyomokat.'
                elif any(kw in desc_lower for kw in ['felirat hiány', 'jelölés hiány']):
                    repair_text = 'Azonosító feliratok, jelölők kialakítása. MSZ HD 60364-6 azonosíthatósági követelmények.'
                else:
                    repair_text = 'A hibához illeszkedő szakszerű javítás végrehajtása. MSZ HD 60364-6:2017 és VMBSZ megfelelés.'
            dp.add_run("Javasolt javítási lépések (MSZ alapján): ").bold = True
            dp.add_run(repair_text + "\n\n")

            dp.add_run("Pontos Helyszín: ").bold = True
            dp.add_run(defect.get('location', 'N/A') + "\n")
            
            # Szabvány hivatkozás - automatikus kitöltés, ha nincs megadva
            dp.add_run("Szabvány hivatkozás: ").bold = True
            std_ref = defect.get('standard', '') or ''
            if not std_ref.strip():
                # Automatikus szabványhivatkozás a hiba tartalom alapján
                if any(kw in desc_lower for kw in ['védővezető', 'rpe', 'folytonosság']):
                    std_ref = 'MSZ HD 60364-6:2017 §61.3.2 (Védővezető folytonosság); 40/2017. (XII.4.) NGM 5.§'
                elif any(kw in desc_lower for kw in ['szigetelés', 'riso', 'insulation']):
                    std_ref = 'MSZ HD 60364-6:2017 §61.3.3 (Szigetelési ellenállás); TvMI 7.7:2026.02.01 §4.3'
                elif any(kw in desc_lower for kw in ['hurok', 'hurokellenállás', 'hurokimpedancia', 'zs']):
                    std_ref = 'MSZ HD 60364-6:2017 §61.3.6 (Hurokimpedancia); MSZ HD 60364-4-41:2017 §411.4'
                elif any(kw in desc_lower for kw in ['rcd', 'ávk', 'fi-relé', 'kioldás']):
                    std_ref = 'MSZ HD 60364-6:2017 §61.3.7 (ÁVK vizsgálat); MSZ EN 61008-1'
                elif any(kw in desc_lower for kw in ['eph', 'potenciál', 'földelés']):
                    std_ref = 'MSZ HD 60364-5-54:2011 §544 (EPH rendszer); MSZ HD 60364-4-41:2017 §411.3.1.2'
                elif any(kw in desc_lower for kw in ['tűz', 'égett', 'ív']):
                    std_ref = 'TvMI 7.7:2026.02.01 (Tűzvédelmi felülvizsgálat); 54/2014. (XII.5.) BM rendelet (OTSZ)'
                elif any(kw in desc_lower for kw in ['selv', 'pelv', 'törpe']):
                    std_ref = 'MSZ HD 60364-4-41:2017 §414 (SELV/PELV)'
                elif any(kw in desc_lower for kw in ['szerszám', 'kéziszerszám']):
                    std_ref = 'MSZ EN 60745-1 (Kéziszerszámok biztonsága)-M4'
                else:
                    std_ref = 'MSZ HD 60364-6:2017 (Villamos berendezések hitelesítése); 40/2017. (XII.4.) NGM rendelet (VMBSZ)'
            dp.add_run(std_ref + "\n")
            
            dp.add_run("Javasolt javítási határidő: ").bold = True
            run = dp.add_run(defect.get('deadline', 'N/A'))
            run.font.color.rgb = RGBColor(255, 0, 0)
            
            # Photo insertion
            photo_data = defect.get('photo')
            if photo_data and photo_data.startswith('data:image'):
                try:
                    # remove data:image/jpeg;base64, prefix
                    b64_str = photo_data.split(',')[1]
                    img_bytes = io.BytesIO(base64.b64decode(b64_str))
                    doc.add_paragraph("Fényképes dokumentáció:").bold = True
                    doc.add_paragraph().add_run().add_picture(img_bytes, width=Cm(12))
                except Exception as e:
                    doc.add_paragraph(f"[Hiba a kép beillesztésekor: {str(e)}]")
    
    section_num += 1

    # ── Bejövő hálózati paraméterek ──
    incoming = getattr(report, 'incoming_phases', None) or {}
    
    if incoming and isinstance(incoming, dict) and any(incoming.get(k) for k in ['l1','l2','l3','systemType','mainFuse']):
        doc.add_heading(f'{section_num}. Bejövő Hálózati Paraméterek', level=1)
        in_table = doc.add_table(rows=1, cols=2)
        in_table.style = 'Table Grid'
        _style_table_header(in_table)
        in_table.rows[0].cells[0].text = 'Paraméter'
        in_table.rows[0].cells[1].text = 'Érték'
        in_rows = [
            ('Rendszer típusa', incoming.get('systemType', 'N/A')),
            ('Fázisszám', f"{incoming.get('phaseCount', '3')} fázis"),
            ('L1 feszültség', f"{incoming.get('l1', 'N/A')} V" if incoming.get('l1') else 'N/A'),
            ('L2 feszültség', f"{incoming.get('l2', 'N/A')} V" if incoming.get('l2') else 'N/A'),
            ('L3 feszültség', f"{incoming.get('l3', 'N/A')} V" if incoming.get('l3') else 'N/A'),
            ('Főbiztosíték (névleges)', f"{incoming.get('mainFuse', 'N/A')} A" if incoming.get('mainFuse') else 'N/A'),
            ('Főbiztosíték (karakterisztika)', incoming.get('mainFuseType', 'N/A') or 'N/A'),
            ('Megjegyzés', incoming.get('note', '') or '—'),
        ]
        for label, val in in_rows:
            rr = in_table.add_row().cells
            rr[0].text = label
            rr[1].text = str(val)
        doc.add_paragraph()
        section_num += 1

    # ── Felülvizsgálói megjegyzés (képekkel) ──
    inspector_notes = getattr(report, 'inspector_notes', '') or ''
    note_photos = getattr(report, 'note_photos', []) or []
    if inspector_notes or note_photos:
        doc.add_page_break()
        doc.add_heading(f'{section_num}. Felülvizsgálói Megjegyzések', level=1)
        if inspector_notes:
            notes_p = doc.add_paragraph(inspector_notes)
            notes_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        if note_photos:
            doc.add_paragraph()
            doc.add_heading(f'{section_num}.1 Csatolt képek', level=2)
            for i, photo_data in enumerate(note_photos, 1):
                if not photo_data or not isinstance(photo_data, str):
                    continue
                try:
                    b64_str = photo_data.split(',')[1] if ',' in photo_data else photo_data
                    img_bytes = io.BytesIO(base64.b64decode(b64_str))
                    photo_p = doc.add_paragraph()
                    photo_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    photo_p.add_run().add_picture(img_bytes, width=Cm(14))
                    cap_p = doc.add_paragraph(f'{i}. kép – Felülvizsgálói melléklet')
                    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in cap_p.runs:
                        r.font.size = Pt(9)
                        r.font.color.rgb = RGBColor(107, 114, 128)
                        r.italic = True
                except Exception as e:
                    logger.warning('Hiba a megjegyzés kép beillesztésekor: %s', e)
                    doc.add_paragraph(f'[Megjegyzés kép #{i}: {str(e)}]')
        section_num += 1

    # ── Alaprajz melléklet ──
    floor_plan_b64 = getattr(report, 'floor_plan_image', None)
    if floor_plan_b64 and isinstance(floor_plan_b64, str) and floor_plan_b64.startswith('data:image'):
        doc.add_page_break()
        doc.add_heading(f'{section_num}. Alaprajz Melléklet', level=1)
        doc.add_paragraph(
            'Az alábbi alaprajz a vizsgált épület / helyszín elrendezését mutatja be, '
            'amelyre az egyvonalas rajz és a mérési eredmények vonatkoznak.'
        )
        try:
            b64_str = floor_plan_b64.split(',')[1] if ',' in floor_plan_b64 else floor_plan_b64
            img_bytes = io.BytesIO(base64.b64decode(b64_str))
            fp_p = doc.add_paragraph()
            fp_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            fp_p.add_run().add_picture(img_bytes, width=Cm(16))
            cap_p = doc.add_paragraph('Alaprajz – vizsgált helyszín elrendezése')
            cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap_p.runs:
                r.font.size = Pt(9)
                r.font.color.rgb = RGBColor(107, 114, 128)
                r.italic = True
        except Exception as e:
            logger.warning('Hiba az alaprajz beillesztésekor: %s', e)
            doc.add_paragraph(f'[Alaprajz beillesztési hiba: {str(e)}]')
        section_num += 1
            
    # Result - Minősítő Irat Változat kezelése
    doc.add_heading(f'{section_num}. Összefoglaló Minősítés', level=1)


    res_p = doc.add_paragraph()
    r_val = c_data.get('reportResult', c_data.get('autoQualification', 'N/A'))
    
    # Minősítő Irat változat szerinti leírás és szín
    qual_descriptions = {
        'MEGFELELŐ': ('MEGFELELŐ', 'A vizsgált villamos berendezés / rendszer az érvényes szabványoknak és előírásoknak megfelel. Hibát nem tártunk fel. Az üzemeltetés folytatható.', RGBColor(0, 128, 0)),
        'VÁLTOZAT_C': ('C VÁLTOZAT – MEGFELELŐ (kisebb hibákkal)', 'A vizsgált villamos berendezés az MSZ HD 60364 szerint alapvetően megfelelő, azonban kisebb eltérések / hibák kerültek megállapításra, amelyek azonnali veszélyt nem jelentenek. A hibák kijavítása ajánlott a következő időszakos felülvizsgálatig.', RGBColor(255, 140, 0)),
        'VÁLTOZAT_B': ('B VÁLTOZAT – FELTÉTELESEN MEGFELELŐ', 'Súlyos hiba(k) kerültek feltárásra. A hibák kijavítása kötelező! Az ismételt ellenőrzés a javítás elvégzése után szükséges. A berendezés a hibák kijavításáig csak fokozott felügyelet mellett üzemeltethető.', RGBColor(255, 80, 0)),
        'VÁLTOZAT_A': ('A VÁLTOZAT – PÓTLÓLAGOS ELLENŐRZÉS SZÜKSÉGES', 'Az érintésvédelmi berendezés a feltárt hibák kijavítása után ismételt, teljes körű ellenőrzésre szorul. A berendezés üzembiztos állapotba hozása és az ismételt felülvizsgálat elvégzése a tulajdonos / üzemeltető felelőssége.', RGBColor(200, 0, 0)),
        'NEM MEGFELELŐ': ('NEM MEGFELELŐ – KÖZVETLEN VESZÉLY', 'A vizsgált villamos berendezés közvetlen élet- és/vagy tűzveszélyes állapotban van! Az azonnali üzemen kívül helyezés és a szakszerű javítás KÖTELEZŐ! A berendezés további üzemeltetése tilos.', RGBColor(255, 0, 0)),
    }
    
    title_text, desc_text, color = qual_descriptions.get(r_val, (r_val, '', RGBColor(128, 128, 128)))
    
    res_run = res_p.add_run(title_text)
    res_run.bold = True
    res_run.font.size = Pt(14)
    res_run.font.color.rgb = color
    
    if desc_text:
        res_p.add_run('\n\n')
        desc_run = res_p.add_run(desc_text)
        desc_run.font.size = Pt(11)
        
    disclaimer = (
        "Módszertan és Felelősség:\n"
        "Az ellenőrzés csak a szemrevételezéssel és jelentős bontás nélkül megközelíthető részekre, "
        "valamint a vonatkozó szabványban (MSZ HD 60364-6) rögzített tesztpontokon végzett műszeres mérésekre terjedt ki. "
        "A rejtett, falban vagy vezetékcsatornában húzódó, elburkolt szerelvények és vezetékek állapota, azok esetleges "
        "szemmel láthatatlan vagy roncsolás nélkül nem vizsgálható hibái a jelenlegi felülvizsgálat során teljes bizonyossággal nem állapíthatók meg.\n\n"
        "A minősítés és a feltárt hibák dokumentálása a felülvizsgálat időpontjában érvényes jogszabályok "
        "alapján történt. A feltárt hibák szakszerű, villamosipari szakember "
        "által történő kijavításáról, valamint a hálózat folyamatos, rendeltetésszerű karbantartásáról az Üzemeltető / Tulajdonos "
        "köteles gondoskodni a VMBSZ rendeletei értelmében.\n\n"
        "Jelen jegyzőkönyv a kiállítás dátumától számított, a vonatkozó rendelet szerinti időszakos "
        "felülvizsgálati határnapig (vagy a berendezés, hálózat átalakításáig) érvényes. "
        "A jegyzőkönyv papír alapú aláírás és a felülvizsgáló bélyegzője nélkül is teljes bizonyító erejű elektronikus okiratnak minősül, amennyiben a "
        "visszakövethető informatikai rendszerből, lezárt és véglegesített (titkosított) formában kerül kiadásra, generálásra."
    )
    p_disc = doc.add_paragraph(disclaimer)
    p_disc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    
    # Word formátum (tervezet / szerkeszthető) nem tartalmaz SHA-256 időbélyeget és QR kódot.
    
    # Következő felülvizsgálat dátuma (OTSZ alapján)
    otsz_class = c_data.get('buildingOtsz', '')
    otsz_cycles = {'AK': 6, 'KK': 3, 'MK': 1, 'NAK': 6}
    if otsz_class and otsz_class in otsz_cycles:
        years = otsz_cycles[otsz_class]
        next_date = datetime.now().replace(year=datetime.now().year + years)
        next_p = doc.add_paragraph()
        next_p.add_run(f"\n📅 KÖVETKEZŐ KÖTELEZŐ FELÜLVIZSGÁLAT: ").bold = True
        next_run = next_p.add_run(f"{next_date.strftime('%Y-%m-%d')}")
        next_run.bold = True
        next_run.font.size = Pt(14)
        next_run.font.color.rgb = RGBColor(255, 80, 0)
        next_p.add_run(f"\n(OTSZ osztály: {otsz_class} → {years} évente, 54/2014. BM rendelet)")
    
    # Aláírás: feltöltött kép (ha van), különben szöveges placeholder
    sig = doc.add_paragraph("\n\n..................................................\nAláírás és Bélyegző")
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if settings and getattr(settings, 'signature_path', None) and os.path.exists(settings.signature_path):
        try:
            from PIL import Image as PilImage
            import io as _io
            sig_img = PilImage.open(settings.signature_path)
            sig_buf = _io.BytesIO()
            sig_img.save(sig_buf, format='PNG')
            sig_buf.seek(0)
            sig_para = doc.add_paragraph()
            sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            sig_run = sig_para.add_run()
            sig_run.add_picture(sig_buf, width=Cm(4.5))
        except Exception as e:
            print(f"Aláírás kép beillesztési hiba: {e}")
    
    # --- Függelék: Villanyszerelői Javítási Nyilatkozat (Opcionális panel) ---
    if "NEM MEGFELELŐ" in r_val or "FELTÉTELESEN" in r_val:
        doc.add_page_break()
        doc.add_heading('Függelék: Villanyszerelői Nyilatkozat a Javításról', level=1)
        
        rep_id_str_f = f"{rep_type}-{report.id}" if report.id else "TERVEZET"
        nyilatkozat_szoveg = (
            f"Alulírott ..................................................... (kivitelező/villanyszerelő neve/cége), "
            "mint megfelelő szakmai képesítéssel rendelkező villamos szakember (Fnyv. szám / bizonyítvány: ................................) "
            f"büntetőjogi felelősségem tudatában nyilatkozom, hogy a jelen, {rep_id_str_f} hivatkozási számú jegyzőkönyvben ("
            "vagy annak hibajegyzék mellékletében) rögzített feltárt hibákat és hiányosságokat a vonatkozó MSZ HD 60364 szabványsorozat előírásainak megfelelően "
            "szakszerűen, maradéktalanul kijavítottam.\n\n"
            "Kijelentem, hogy az általam elvégzett javítási munkálatok után a vizsgált villamos berendezés/hálózat áramütés elleni védelme, "
            "valamint szabványos állapota a biztonságos üzemeltetés feltételeinek maradéktalanul megfelel.\n\n"
        )
        j_p = doc.add_paragraph(nyilatkozat_szoveg)
        j_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph("Javítás elvégzésének dátuma: 20... év ...................... hó ........ nap\n\n\n\n")
        
        j_sig = doc.add_paragraph("..................................................\nJavítást Végző Villanyszerelő aláírása")
        j_sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        if settings and getattr(settings, 'signature_path', None) and os.path.exists(settings.signature_path):
            try:
                from PIL import Image as PilImage
                import io as _io
                sig_img = PilImage.open(settings.signature_path)
                sig_buf = _io.BytesIO()
                sig_img.save(sig_buf, format='PNG')
                sig_buf.seek(0)
                j_sig_para = doc.add_paragraph()
                j_sig_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                j_sig_para.add_run().add_picture(sig_buf, width=Cm(4.5))
            except Exception as e:
                print(f"Javítás aláírás kép beillesztési hiba: {e}")

    stream = io.BytesIO()
    doc.save(stream)
    stream.seek(0)
    return stream

import tempfile
import os


def apply_pdf_watermark(pdf_bytes: bytes, watermark_text: Optional[str] = None) -> bytes:
    """Átlátszó, ferde vízjel minden oldalra (demó / ingyenes csomag)."""
    text = (watermark_text or os.environ.get("PDF_WATERMARK_TEXT") or "DEMÓ – VBF Premium").strip()
    from io import BytesIO as _BIO

    from reportlab.pdfgen import canvas
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.colors import Color
    from pypdf import PdfReader, PdfWriter

    w, h = A4
    packet = _BIO()
    c = canvas.Canvas(packet, pagesize=A4)
    c.saveState()
    c.setFillColor(Color(0.5, 0.5, 0.5, alpha=0.22))
    c.setFont("Helvetica-Bold", 34)
    c.translate(w / 2, h / 2)
    c.rotate(32)
    c.drawCentredString(0, 0, text)
    c.restoreState()
    c.save()
    packet.seek(0)
    wm_reader = PdfReader(packet)
    wm_page = wm_reader.pages[0]

    reader = PdfReader(_BIO(pdf_bytes))
    writer = PdfWriter()
    for page in reader.pages:
        page.merge_page(wm_page)
        writer.add_page(page)
    out = _BIO()
    writer.write(out)
    return out.getvalue()


def generate_signed_pdf_stream(
    report: Report,
    db=None,
    pfx_path: Optional[str] = None,
    pfx_pass: Optional[bytes] = None,
    share_url: Optional[str] = None,
    watermark: bool = False,
) -> io.BytesIO:
    # PFX: először megadott path/jelszó, majd céges beállítások, végül alapértelmezett
    if pfx_path is None and db and report.owner_id:
        owner = db.query(database.User).filter(database.User.id == report.owner_id).first()
        if owner and getattr(owner, "company_id", None):
            settings = db.query(database.CompanySettings).filter(
                database.CompanySettings.company_id == owner.company_id
            ).first()
        else:
            settings = db.query(database.CompanySettings).filter(
                database.CompanySettings.owner_id == report.owner_id
            ).first()
        if settings and getattr(settings, "pfx_path", None) and os.path.exists(settings.pfx_path):
            pfx_path = settings.pfx_path
    if pfx_path is None:
        pfx_path = "signer.pfx"
    if pfx_pass is None:
        import os as _os
        # Biztonság: ha nincs jelszó beállítva, ne használjunk gyenge defaultot,
        # inkább próbáljuk meg aláírás nélkül visszaadni a PDF-et.
        pfx_pass = (_os.environ.get("VBF_PFX_PASSWORD") or "").encode("utf-8")

    # ReportLab-bal közvetlen PDF – nincs LibreOffice, alacsony erőforrás-igény
    from generator_pdf import generate_pdf_reportlab_stream
    pdf_buffer = generate_pdf_reportlab_stream(report, db, share_url=share_url)
    pdf_bytes = pdf_buffer.getvalue()

    if watermark:
        pdf_bytes = apply_pdf_watermark(pdf_bytes)
        out_plain = io.BytesIO(pdf_bytes)
        out_plain.seek(0)
        return out_plain

    from pyhanko.sign import signers
    from pyhanko.pdf_utils.reader import PdfFileReader
    from pyhanko.pdf_utils.writer import copy_into_new_writer
    import io as pyhanko_io

    if not os.path.exists(pfx_path):
        pdf_out = pyhanko_io.BytesIO(pdf_bytes)
        pdf_out.seek(0)
        return pdf_out
    if not pfx_pass:
        pdf_out = pyhanko_io.BytesIO(pdf_bytes)
        pdf_out.seek(0)
        return pdf_out

    signer = signers.SimpleSigner.load_pkcs12(pfx_path, pfx_pass)
    r = PdfFileReader(pyhanko_io.BytesIO(pdf_bytes))
    w = copy_into_new_writer(r)
    pdf_out = pyhanko_io.BytesIO()
    signers.sign_pdf(
        w,
        signers.PdfSignatureMetadata(field_name='VBF_Signature'),
        signer=signer,
        in_place=False,
        existing_fields_only=False,
        bytes_reserved=8192,
        out=pdf_out
    )
    pdf_out.seek(0)
    return pdf_out
